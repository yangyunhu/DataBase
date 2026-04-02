#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word转PDF脚本 - 使用多种方案确保表格格式正确转换
"""
import sys
import os
import subprocess
import time

def convert_with_libreoffice(docx_path, output_path):
    """使用LibreOffice转换 - 禁用弹窗"""
    try:
        print("尝试使用LibreOffice转换...")
        
        # 查找LibreOffice路径
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\ProgramData\chocolatey\lib\libreoffice\tools\LibreOffice\program\soffice.exe",
        ]
        
        soffice_path = None
        for path in libreoffice_paths:
            if os.path.exists(path):
                soffice_path = path
                break
        
        if not soffice_path:
            print("LibreOffice未找到，跳过")
            return False
        
        output_dir = os.path.dirname(output_path)
        
        # 构建命令 - 使用更严格的参数避免弹窗
        cmd = [
            soffice_path,
            '--headless',
            '--nologo',
            '--nodefault',
            '--nofirststartwizard',
            '--norestore',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        
        print(f"执行命令: {' '.join(cmd)}")
        
        # 执行转换，完全隐藏窗口
        try:
            # 使用STARTUPINFO隐藏窗口
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            startupinfo.wShowWindow = subprocess.SW_HIDE
            
            result = subprocess.run(
                cmd, 
                capture_output=True, 
                text=True, 
                timeout=30,
                startupinfo=startupinfo,
                creationflags=subprocess.CREATE_NO_WINDOW | subprocess.CREATE_NEW_PROCESS_GROUP
            )
            
            # 检查是否有错误输出
            if result.returncode != 0:
                print(f"LibreOffice返回错误码: {result.returncode}")
                print(f"LibreOffice错误: {result.stderr}")
                return False
            
        except subprocess.TimeoutExpired:
            print("LibreOffice转换超时，跳过")
            return False
        except Exception as e:
            print(f"LibreOffice执行失败: {e}")
            return False
        
        # LibreOffice生成的文件名
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        generated_pdf = os.path.join(output_dir, base_name + '.pdf')
        
        # 等待文件生成
        max_wait = 10
        waited = 0
        while waited < max_wait and not os.path.exists(generated_pdf):
            time.sleep(0.5)
            waited += 0.5
        
        if os.path.exists(generated_pdf):
            # 重命名为目标文件名
            if generated_pdf != output_path:
                os.rename(generated_pdf, output_path)
            print(f"LibreOffice转换成功: {output_path}")
            return True
        else:
            print("LibreOffice转换失败: 输出文件未生成")
            return False
            
    except Exception as e:
        print(f"LibreOffice转换失败: {e}")
        return False

def convert_with_docx2pdf(docx_path, output_path):
    """使用docx2pdf转换"""
    try:
        print("尝试使用docx2pdf转换...")
        from docx2pdf import convert
        
        # Windows不支持signal.SIGALRM，使用线程超时替代
        import threading
        
        conversion_result = {"success": False, "completed": False}
        
        def do_conversion():
            try:
                convert(docx_path, output_path)
                conversion_result["success"] = True
            except Exception as e:
                print(f"docx2pdf转换内部错误: {e}")
                conversion_result["success"] = False
            finally:
                conversion_result["completed"] = True
        
        # 启动转换线程
        thread = threading.Thread(target=do_conversion)
        thread.daemon = True
        thread.start()
        
        # 等待最多10秒
        thread.join(timeout=10)
        
        if not conversion_result["completed"]:
            print("docx2pdf转换超时")
            return False
        
        if not conversion_result["success"]:
            return False
        
        # 等待文件生成（最多5秒）
        max_wait = 5
        waited = 0
        while waited < max_wait and not os.path.exists(output_path):
            time.sleep(0.2)
            waited += 0.2
        
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"docx2pdf转换成功: {output_path}")
            return True
        else:
            print("docx2pdf转换失败: 文件未生成")
            return False
            
    except Exception as e:
        print(f"docx2pdf转换失败: {e}")
        return False

def convert_with_pywin32(docx_path, output_path):
    """使用pywin32 COM转换"""
    try:
        print("尝试使用pywin32转换...")
        import win32com.client as win32
        
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        
        # 优化：使用try-finally确保Word进程被关闭
        try:
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(output_path), FileFormat=17)
            doc.Close()
        finally:
            word.Quit()
        
        # 等待文件生成（最多2秒）
        max_wait = 2
        waited = 0
        while waited < max_wait and not os.path.exists(output_path):
            time.sleep(0.1)  # 减少等待间隔
            waited += 0.1
        
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"pywin32转换成功: {output_path}")
            return True
        else:
            print("pywin32转换失败")
            return False
            
    except Exception as e:
        print(f"pywin32转换失败: {e}")
        return False

def convert_with_pypandoc(docx_path, output_path):
    """使用pypandoc转换Word到PDF"""
    try:
        print("尝试使用pypandoc转换...")
        import pypandoc
        
        # 直接转换为PDF，跳过HTML中间步骤
        try:
            # 直接使用pandoc转换
            output = pypandoc.convert_file(
                docx_path,
                'pdf',
                format='docx',
                outputfile=output_path,
                extra_args=['--pdf-engine=wkhtmltopdf']  # 尝试使用wkhtmltopdf
            )
            
            # 等待文件生成（最多3秒）
            max_wait = 3
            waited = 0
            while waited < max_wait and not os.path.exists(output_path):
                time.sleep(0.1)
                waited += 0.1
            
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"pypandoc直接转换成功: {output_path}")
                return True
        except Exception as e:
            print(f"pypandoc直接转换失败: {e}")
        
        # 尝试使用HTML中间步骤
        try:
            # 先将docx转换为html
            html_output = pypandoc.convert_file(docx_path, 'html', format='docx')
            
            # 保存HTML临时文件
            html_path = output_path.replace('.pdf', '.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_output)
            
            # 尝试使用pdfkit（快速）
            try:
                import pdfkit
                # 使用快速模式
                options = {'quiet': ''}
                pdfkit.from_file(html_path, output_path, options=options)
                
                # 清理临时文件
                if os.path.exists(html_path):
                    os.unlink(html_path)
                
                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    print(f"pypandoc+pdfkit转换成功: {output_path}")
                    return True
            except ImportError:
                pass  # 跳过，不打印错误
            except Exception as e:
                pass  # 跳过，不打印错误
            
            # 清理临时文件
            if os.path.exists(html_path):
                os.unlink(html_path)
                
        except Exception as e:
            pass  # 跳过，不打印错误
        
        print("pypandoc转换失败")
        return False
        
    except Exception as e:
        print(f"pypandoc转换失败: {e}")
        return False

def convert_with_reportlab(docx_path, output_path):
    """使用reportlab创建PDF（后备方案）- 支持表格绘制"""
    try:
        print("尝试使用reportlab创建PDF...")
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Table, TableStyle, Paragraph, SimpleDocTemplate, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        
        # 读取Word文档
        docx = Document(docx_path)
        
        # 创建PDF文档
        doc = SimpleDocTemplate(output_path, pagesize=A4, 
                               leftMargin=1.0*cm, rightMargin=1.0*cm, 
                               topMargin=1.0*cm, bottomMargin=1.0*cm)
        
        # 获取样式
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        
        # 尝试注册中文字体
        font_name = "Helvetica"
        try:
            font_paths = [
                r"C:\Windows\Fonts\simsun.ttc",
                r"C:\Windows\Fonts\msyh.ttc",
                r"C:\Windows\Fonts\simhei.ttf",
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        font_name = 'ChineseFont'
                        normal_style.fontName = font_name
                        print(f"使用中文字体: {font_path}")
                        break
                    except:
                        continue
        except Exception as font_error:
            print(f"字体设置失败: {font_error}")
        
        # 构建内容
        story = []
        
        # 处理段落
        for para in docx.paragraphs:
            if para.text.strip():
                p = Paragraph(para.text, normal_style)
                story.append(p)
                story.append(Spacer(1, 12))
        
        # 处理表格
        for table in docx.tables:
            # 提取表格数据
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # 为表格单元格创建Paragraph
                    cell_paragraph = Paragraph(cell_text, normal_style)
                    row_data.append(cell_paragraph)
                data.append(row_data)
            
            if data:
                # 创建表格
                table_width = A4[0] - 2.0*cm  # 表格宽度
                num_cols = len(data[0]) if data else 1
                col_widths = [table_width / num_cols] * num_cols
                
                table = Table(data, colWidths=col_widths)
                
                # 设置表格样式
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), font_name),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                
                story.append(table)
                story.append(Spacer(1, 24))
        
        # 构建PDF
        doc.build(story)
        
        print(f"reportlab PDF创建成功，包含表格")
        
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            file_size = os.path.getsize(output_path)
            print(f"使用 reportlab 转换成功，文件大小: {file_size} 字节")
            return True
        else:
            print("reportlab转换失败")
            return False
            
    except Exception as e:
        print(f"reportlab 转换失败: {e}")
        import traceback
        traceback.print_exc()
        # 回退到纯文本方案
        try:
            print("回退到纯文本方案...")
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import cm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # 读取Word文档
            doc = Document(docx_path)
            
            # 提取所有文本
            all_text = []
            for para in doc.paragraphs:
                all_text.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                all_text.append("--- 表格 ---")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    all_text.append(" | ".join(row_text))
            
            # 创建PDF
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            # 尝试注册中文字体
            font_name = "Helvetica"
            try:
                font_paths = [
                    r"C:\Windows\Fonts\simsun.ttc",
                    r"C:\Windows\Fonts\msyh.ttc",
                    r"C:\Windows\Fonts\simhei.ttf",
                ]
                
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            font_name = 'ChineseFont'
                            break
                        except:
                            continue
            except Exception as font_error:
                print(f"字体设置失败: {font_error}")
            
            c.setFont(font_name, 12)
            
            # 绘制内容
            left_margin = 0.5 * cm
            right_margin = width - 1.5 * cm
            top_margin = height - 1.5 * cm
            bottom_margin = 1.5 * cm
            usable_width = right_margin - left_margin
            line_height = 14
            y = top_margin
            
            for line in all_text:
                if y < bottom_margin:
                    c.showPage()
                    y = top_margin
                    c.setFont(font_name, 12)
                
                if line.strip():
                    text = line.strip()
                    from reportlab.pdfbase.pdfmetrics import stringWidth
                    words = text.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + " " + word if current_line else word
                        text_width = stringWidth(test_line, font_name, 12)
                        
                        if text_width > usable_width:
                            if current_line:
                                c.drawString(left_margin, y, current_line)
                                y -= line_height
                                if y < bottom_margin:
                                    c.showPage()
                                    y = top_margin
                                    c.setFont(font_name, 12)
                            current_line = word
                        else:
                            current_line = test_line
                    
                    if current_line:
                        c.drawString(left_margin, y, current_line)
                        y -= line_height
                else:
                    y -= line_height / 2
            
            c.save()
            print("纯文本方案PDF创建成功")
            
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                return True
            else:
                return False
                
        except Exception as e:
            print(f"回退方案也失败: {e}")
            return False

def convert_word_to_pdf(docx_path, output_path):
    """
    将Word文档转换为PDF，使用多种方案确保表格格式正确
    """
    print(f"开始转换: {docx_path}")
    print(f"输出文件: {output_path}")
    
    # 方案1: docx2pdf
    if convert_with_docx2pdf(docx_path, output_path):
        return True
    
    # 方案2: pywin32
    if convert_with_pywin32(docx_path, output_path):
        return True
    
    # 方案3: pypandoc（使用pandoc转换，保留表格格式）
    if convert_with_pypandoc(docx_path, output_path):
        return True
    
    # 方案4: reportlab（后备方案，支持表格绘制）
    if convert_with_reportlab(docx_path, output_path):
        return True
    
    print("所有转换方案都失败")
    return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python word_to_pdf.py <输入Word文件> <输出PDF文件>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not os.path.exists(docx_path):
        print(f"错误: 文件不存在 {docx_path}")
        sys.exit(1)
    
    success = convert_word_to_pdf(docx_path, output_path)
    sys.exit(0 if success else 1)
