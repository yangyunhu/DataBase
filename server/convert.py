#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档转换脚本
支持 PDF 转 Word 和 Word 转 PDF
"""

import sys
import os
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert as docx_to_pdf

def unify_fonts(docx_path):
    """
    统一 Word 文档中的字体为宋体
    
    Args:
        docx_path: Word 文档路径
    """
    try:
        print(f"统一字体: {docx_path}")
        
        # 打开文档
        doc = Document(docx_path)
        
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            # 遍历段落中的所有 run（文本片段）
            for run in paragraph.runs:
                # 统一设置为宋体
                run.font.name = 'SimSun'
                # 对于中文字体，还需要设置 eastAsia 属性
                run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/drawingml/2006/main}eastAsia', 'SimSun')
        
        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'SimSun'
                            run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/drawingml/2006/main}eastAsia', 'SimSun')
        
        # 保存文档
        doc.save(docx_path)
        print(f"字体统一完成: {docx_path}")
        return True
        
    except Exception as e:
        print(f"字体统一失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def convert_pdf_to_word(pdf_path, output_path):
    """
    将 PDF 文件转换为 Word 文档（优化速度）
    
    Args:
        pdf_path: 输入 PDF 文件路径
        output_path: 输出 Word 文件路径
    """
    try:
        print(f"开始转换: {pdf_path}")
        print(f"输出文件: {output_path}")
        
        # 创建转换器
        cv = Converter(pdf_path)
        
        # 转换 PDF 到 Word - 优化参数
        cv.convert(
            output_path,
            start=0,
            end=None,
            pages=None,
            multi_processing=True,  # 启用多进程转换
            line_margin_ratio=0.05,  # 减少行边距检测时间
            word_margin_ratio=0.02,  # 减少单词边距检测时间
        )
        
        # 关闭转换器
        cv.close()
        
        print(f"PDF 转换完成: {output_path}")
        
        # 跳过字体统一步骤，提高速度
        # unify_fonts(output_path)
        
        print(f"转换成功: {output_path}")
        return True
        
    except Exception as e:
        print(f"转换失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def convert_word_to_pdf(docx_path, output_path):
    """
    将 Word 文档转换为 PDF，使用多种方案确保表格格式正确
    
    Args:
        docx_path: 输入 Word 文档路径
        output_path: 输出 PDF 文件路径
    """
    try:
        print(f"开始转换: {docx_path}")
        print(f"输出文件: {output_path}")
        
        # 导入word_to_pdf模块
        import importlib.util
        word_to_pdf_path = os.path.join(os.path.dirname(__file__), 'word_to_pdf.py')
        
        spec = importlib.util.spec_from_file_location("word_to_pdf", word_to_pdf_path)
        word_to_pdf_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(word_to_pdf_module)
        
        # 使用新的转换函数
        success = word_to_pdf_module.convert_word_to_pdf(docx_path, output_path)
        
        if success:
            # 验证PDF文件
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                with open(output_path, 'rb') as f:
                    header = f.read(4)
                    if header == b'%PDF':
                        file_size = os.path.getsize(output_path)
                        print(f"Word 转换完成: {output_path}")
                        print(f"转换成功，文件大小: {file_size} 字节")
                        return True
                    else:
                        print(f"PDF文件头不正确: {header}")
                        return False
            else:
                print("PDF文件未生成或为空")
                return False
        else:
            print("转换失败: 所有方案都失败")
            return False
        
    except Exception as e:
        print(f"转换失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

import json

if __name__ == "__main__":
    if len(sys.argv) != 4:
        result = {
            "success": False,
            "error": "参数错误: python convert.py <转换类型> <输入文件路径> <输出文件路径>"
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
    
    conversion_type = sys.argv[1]
    input_path = sys.argv[2]
    output_path = sys.argv[3]
    
    try:
        if conversion_type == "pdf2word":
            success = convert_pdf_to_word(input_path, output_path)
        elif conversion_type == "word2pdf":
            success = convert_word_to_pdf(input_path, output_path)
        else:
            result = {
                "success": False,
                "error": f"不支持的转换类型: {conversion_type}"
            }
            print(json.dumps(result, ensure_ascii=False))
            sys.exit(1)
        
        if success:
            result = {
                "success": True,
                "output_path": output_path,
                "message": "转换成功"
            }
        else:
            result = {
                "success": False,
                "error": "转换失败"
            }
        
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(0 if success else 1)
        
    except Exception as e:
        result = {
            "success": False,
            "error": str(e)
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
